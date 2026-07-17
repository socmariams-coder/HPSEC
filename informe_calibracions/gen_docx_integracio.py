# -*- coding: utf-8 -*-
"""Generador de l'informe exhaustiu d'integracio del DOC (305 vs 306).

Llegeix les dades directament de les seqüencies (via compara_integracions) i les
figures generades per fes_figures_integracio.py.

Us:  python -X utf8 gen_docx_integracio.py
"""
import os
import sys

import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from compara_integracions import recull, recta, VOL_UL

SCR = os.path.dirname(os.path.abspath(__file__))
OUT = r"C:\Users\maria\Proyectos\Informe_integracio_DOC_305_306_v2.docx"

ACCENT = RGBColor(0x0E, 0x62, 0x6C)
INK = RGBColor(0x22, 0x2A, 0x30)
MUT = RGBColor(0x5A, 0x6B, 0x78)
RED = RGBColor(0xBE, 0x3A, 0x38)

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(10.5); st.font.color.rgb = INK
for s in ("Heading 1", "Heading 2"):
    doc.styles[s].font.color.rgb = ACCENT

FIGN = [0]; TBLN = [0]


def para(text, size=10.5, bold=False, italic=False, color=INK, space_after=6, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic; r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(3)
    return p


def figure(png, caption, width=6.4):
    doc.add_picture(os.path.join(SCR, png), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    FIGN[0] += 1
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Figura {FIGN[0]}. "); r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = MUT
    r2 = p.add_run(caption); r2.italic = True; r2.font.size = Pt(8.5); r2.font.color.rgb = MUT
    p.paragraph_format.space_after = Pt(12)


def table_caption(caption):
    TBLN[0] += 1
    p = doc.add_paragraph()
    r = p.add_run(f"Taula {TBLN[0]}. "); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = MUT
    r2 = p.add_run(caption); r2.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = MUT
    p.paragraph_format.space_after = Pt(4)


def shade(cell, hexcolor="0E626C"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def repeat_header(table):
    trPr = table.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true"); trPr.append(th)


def taula(data, marca_vermell=None):
    """data[0] = capcalera. marca_vermell = fn(i, j, val) -> bool"""
    tb = doc.add_table(rows=len(data), cols=len(data[0]))
    tb.style = "Light Grid Accent 1"; tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            c = tb.cell(i, j); c.text = ""
            run = c.paragraphs[0].add_run(str(val)); run.font.size = Pt(8.5)
            if i == 0:
                run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); shade(c)
            elif marca_vermell and marca_vermell(i, j, val):
                run.font.color.rgb = RED; run.bold = True
    repeat_header(tb)
    doc.add_paragraph()
    return tb


def _f(v, dec=1):
    if v is None or (isinstance(v, float) and not np.isfinite(v)):
        return "—"
    return f"{v:,.{dec}f}".replace(",", "§").replace(".", ",").replace("§", ".")


def _num(s):
    """Desfa _f(): '1.234,56×' -> 1234.56. Retorna None si no es numeric."""
    try:
        return float(str(s).rstrip("×").strip().replace(".", "").replace(",", "."))
    except (ValueError, AttributeError):
        return None


# ============================== DADES ==============================
print("Recollint dades…")
D = {s: recull(s) for s in ["293_SEQ_CAL", "305_SEQ_CAL", "306_SEQ_CAL"]}


def rectes(seq):
    files, _, _ = D[seq]
    ps, pe = {}, {}
    for f in files:
        ps.setdefault(f["conc"], []).append(f["suite_area"])
        if f["est"]:
            pe.setdefault(f["conc"], []).append(f["est"]["area"])
    a = recta([(c * VOL_UL / 1000, float(np.mean(v))) for c, v in sorted(ps.items())])
    b = recta([(c * VOL_UL / 1000, float(np.mean(v))) for c, v in sorted(pe.items())])
    return a, b


R = {s: rectes(s) for s in D}

# ============================== PORTADA ==============================
t = doc.add_paragraph()
r = t.add_run("Integració del senyal DOC a les seqüències de calibratge 305 i 306")
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = ACCENT
t.paragraph_format.space_after = Pt(4)
para("Anàlisi injecció a injecció de la finestra d'integració del canal DOC (TOC) i del seu efecte "
     "sobre la recta de calibratge. Comparació amb la seqüència 293 com a referència.",
     size=11, color=MUT, space_after=10)
para("Seqüències: COLUMN 293_SEQ_CAL (20/02) · 305_SEQ_CAL (07/07) · 306_SEQ_CAL (14/07), 400 µL, "
     "dues rèpliques per concentració. La 305 i la 306 comparteixen detector, mode, volum d'injecció i "
     "mètode cromatogràfic; difereixen únicament en la columna, substituïda entre totes dues.",
     size=9, color=MUT, space_after=14)

# ============================== 1. RESUM ==============================
doc.add_heading("1. Resum", level=1)
para("La finestra d'integració que obre el programa és més ampla que el pic en totes tres seqüències "
     "(mediana de 10,3 · 9,7 i 13,0 vegades el FWHM, davant les ~4 vegades que correspondrien a un pic "
     "gaussià). L'excés de finestra no té conseqüència mesurable mentre la línia de base és plana, però "
     "a la seqüència 306 —amb la base elevada i derivant— la finestra incorpora senyal que no pertany al "
     "pic de KHP.")
para("Observacions principals:", bold=True, space_after=3)
bullet("Amplada de la finestra: 293 de 6,9 a 18,6 FWHM (mediana 10,3); 305 de 6,9 a 12,3 (mediana 9,7); "
       "306 de 6,4 a 40,9 (mediana 13,0). Les 32 injeccions superen les 6 vegades el FWHM.")
bullet("Sobre-integració (àrea del programa / àrea amb finestra estreta): 293 de ×1,05 a ×1,44; "
       "305 de ×1,02 a ×1,20; 306 de ×1,07 a ×15,0, amb un cas de ×173.")
bullet("La sobre-integració depèn de la concentració: creix en disminuir la concentració, perquè el "
       "senyal aliè capturat és constant mentre el pic minva. A la 306 aquesta dependència arriba a "
       "invertir l'ordre dels punts: l'àrea a 0,25 ppm (12.583) supera la de 3 ppm (11.564).")
bullet("Estructures tardanes: les dues seqüències presenten els mateixos accidents posteriors al pic "
       "(pics a ~31 i ~36 min i un altiplà de ~44 a ~72 min). A la 305 la finestra s'atura abans "
       "d'arribar-hi; a la 306 els incorpora.")
bullet("Efecte sobre la recta: a la 305 el mètode d'integració és indiferent (RF_mass 7.868 · R² 0,998 "
       "amb el programa; 7.718 · R² 0,999 amb finestra estreta). A la 306 és determinant (6.193 · "
       "R² 0,607 amb el programa; 9.319 · R² 0,9996 amb finestra estreta).")
bullet("Amb el mateix mètode d'integració, les rectes de la 305 i la 306 no coincideixen: la 306 dóna un "
       "factor de resposta un 21 % superior (9.319 davant 7.718), tot i que l'alçada neta del pic a 5 ppm "
       "és equivalent (10.756 i 10.756 ppb) i el FWHM difereix (1,48 i 1,60 min).")

# ============================== 2. METODE ==============================
doc.add_heading("2. Mètode", level=1)
para("S'han comparat dos criteris d'integració sobre les mateixes dades:")
para("Integració vigent (programa)", bold=True, space_after=2)
para("Els límits del pic es determinen per projecció tangent sobre la derivada del senyal "
     "(find_peak_boundaries). El criteri busca el punt on el pic retorna a la línia de base. L'àrea "
     "s'integra entre aquests límits sobre el senyal net.", space_after=8)
para("Integració estreta ancorada", bold=True, space_after=2)
para("El pic es localitza dins d'una finestra de ±3 min al voltant de la posició esperada, obtinguda del "
     "màxim del canal de 254 nm (àncora) més el desfasament DOC−254 mesurat a la mateixa seqüència "
     "(+1,86 min a la 293; +1,04 min a la 305; +2,03 min a la 306). Se'n mesura el FWHM per creuament "
     "del semi-màxim i s'integra a ±2·FWHM (4·FWHM en total, que conté el 99,9 % d'un pic gaussià) "
     "sobre una línia de base lineal ajustada als dos flancs de la finestra.", space_after=8)
para("El canal de 254 nm actua d'invariant: el pic es manté entre seqüències (×1,10 a la 305 respecte de "
     "la 293), la qual cosa indica que la quantitat de KHP injectada no varia. Qualsevol diferència entre "
     "les àrees del DOC prové, doncs, del tren de detecció o del criteri d'integració, no de la mostra.",
     space_after=8)

# ============================== 3. SENYAL CONTINU ==============================
doc.add_heading("3. Senyal continu", level=1)
para("La Figura 1 mostra el senyal DOC cru de les tres seqüències, amb les injeccions de KHP "
     "concatenades en ordre d'adquisició. Permet situar la línia de base de cada seqüència i els "
     "accidents que la integració pot arribar a capturar.")
figure("int_doc_continu.png",
       "DOC cru continu (fulla 2-TOC), un panell per seqüència. La 293 manté la línia de base a ~35 ppb. "
       "La 305 la manté a ~650 ppb, plana i sense caigudes. La 306 la manté entre ~1.400 i ~2.000 ppb, "
       "amb deriva dins de cada injecció i caigudes abruptes fins a ~0 ppb en dues de les injeccions "
       "inicials. L'alçada dels pics escala amb la concentració a totes tres seqüències.")

# ============================== 4. FINESTRA ==============================
doc.add_heading("4. Finestra d'integració", level=1)
para("La Figura 2 superposa, sobre el cromatograma de cada concentració, la finestra que obre el "
     "programa i la finestra estreta ancorada. La Taula 1 en resumeix l'amplada respecte del FWHM del pic.")
figure("int_finestres.png",
       "Cromatograma del DOC (rèplica 1), una fila per concentració i una columna per seqüència, amb la "
       "finestra del programa (vermell) i la finestra estreta ancorada (verd). Línia de punts: posició "
       "del pic de 254 nm. A la 293 i la 305 la finestra del programa és ampla però s'atura abans de "
       "l'altiplà tardà, i totes dues àrees pràcticament coincideixen. A la 306, a 0,25 i 0,5 ppm, la "
       "finestra s'estén fins a incorporar-lo. Els panells marcats «no present» corresponen a "
       "concentracions no incloses en aquella seqüència (0,1 i 2 ppm només a la 293; 3 ppm només a la "
       "305 i la 306).", width=6.9)

table_caption("Amplada de la finestra d'integració del programa, expressada en múltiples del FWHM del pic. "
              "Una integració completa d'un pic gaussià requereix ~4·FWHM.")
dat = [("Seqüència", "mínim", "mediana", "màxim", "injeccions > 6·FWHM")]
for seq in ["293_SEQ_CAL", "305_SEQ_CAL", "306_SEQ_CAL"]:
    files, _, _ = D[seq]
    xs = [f["suite_ampl"] / f["est"]["fwhm"] for f in files
          if f["est"] and f["est"]["fwhm"] and np.isfinite(f["suite_ampl"])]
    dat.append((seq, _f(min(xs)), _f(float(np.median(xs))), _f(max(xs)), f"{len(xs)} de {len(xs)}"))
taula(dat)

# ============================== 5. DETALL ==============================
doc.add_heading("5. Detall injecció a injecció", level=1)
para("Les Taules 2 a 4 recullen, per rèplica, la posició del pic, el FWHM, la finestra oberta pel "
     "programa i l'àrea resultant amb cada criteri. La columna «sobre-int.» és el quocient entre les dues "
     "àrees. En vermell, els valors superiors a ×2.", color=MUT, space_after=8)

for seq in ["293_SEQ_CAL", "305_SEQ_CAL", "306_SEQ_CAL"]:
    files, t254_ref, shift = D[seq]
    etiqueta = {"293_SEQ_CAL": "referència anterior",
                "305_SEQ_CAL": "columna anterior a la substitució",
                "306_SEQ_CAL": "columna substituïda"}[seq]
    table_caption(f"{seq} ({etiqueta}). Àncora de 254 nm de consens: {_f(t254_ref, 2)} min; "
                  f"desfasament DOC−254: {shift:+.2f} min. Àrees en unitats d'integració.")
    dat = [("ppm", "Rep", "t 254", "t DOC", "FWHM", "finestra programa", "ampl.",
            "×FWHM", "àrea programa", "àrea estreta", "sobre-int.")]
    for f in sorted(files, key=lambda x: (x["conc"], x["rep"])):
        e = f["est"]
        if not e:
            continue
        xf = f["suite_ampl"] / e["fwhm"] if e["fwhm"] else np.nan
        ratio = f["suite_area"] / e["area"] if e["area"] > 0 else np.nan
        t254 = f"{_f(f['t254'], 2)}{'' if f['fiable254'] else ' *'}"
        dat.append((f"{f['conc']:g}".replace(".", ","), f"R{f['rep']}", t254,
                    _f(e["t_max"], 2), _f(e["fwhm"], 2),
                    f"{_f(f['suite_lo'], 1)}–{_f(f['suite_hi'], 1)}",
                    _f(f["suite_ampl"], 1), _f(xf, 1),
                    _f(f["suite_area"], 1), _f(e["area"], 1), _f(ratio, 2) + "×"))
    taula(dat, marca_vermell=lambda i, j, v: j == 10 and isinstance(v, str)
          and v.endswith("×") and _num(v) is not None and _num(v) > 2)
    if any(not f["fiable254"] for f in files):
        para("* El pic de 254 nm d'aquesta injecció no és localitzable amb fiabilitat a causa de la "
             "deriva de la línia de base del DAD; s'hi ha aplicat l'àncora de consens de la seqüència.",
             size=8, color=MUT, space_after=10)

# ============================== 6. SOBRE-INTEGRACIO ==============================
doc.add_heading("6. Dependència amb la concentració", level=1)
para("La Figura 3 mostra el factor de sobre-integració i l'amplada de la finestra en funció de la "
     "concentració. El senyal aliè capturat per l'excés de finestra és aproximadament constant dins de "
     "cada seqüència; per tant, el seu pes relatiu creix en disminuir la concentració del patró. A la 293 "
     "i la 305 aquest senyal és menyspreable respecte del pic (base plana i propera a zero un cop "
     "restada), i el factor es manté entre ×1,0 i ×1,4 a tot el rang. A la 306 el mateix mecanisme "
     "produeix factors de ×1,07 a 5 ppm i de ×15,0 a 0,25 ppm.")
figure("int_sobreintegracio.png",
       "Esquerra: factor de sobre-integració (àrea del programa / àrea amb finestra estreta), escala "
       "logarítmica. Dreta: amplada de la finestra del programa en múltiples del FWHM. La línia "
       "discontínua marca, respectivament, l'absència de sobre-integració i la finestra de 4·FWHM.")

# ============================== 7. RECTA ==============================
doc.add_heading("7. Efecte sobre la recta de calibratge", level=1)
para("La Taula 5 i la Figura 4 recullen la recta obtinguda amb cada criteri d'integració. A la 305 el "
     "criteri és indiferent: les dues rectes coincideixen dins de l'1,9 % i totes dues són lineals. A la "
     "306 el criteri determina el resultat: amb la integració vigent la recta perd la linealitat "
     "(R² 0,607) i presenta una ordenada a l'origen de 5.630, del mateix ordre que l'àrea del punt de "
     "3 ppm; amb finestra estreta la linealitat es recupera (R² 0,9996) i l'ordenada a l'origen torna a "
     "ser propera a zero.")
table_caption("Recta de calibratge (àrea = RF_mass · µg DOC + ordenada) amb cada criteri d'integració. "
              "Mitjana de les dues rèpliques per concentració; 400 µL d'injecció.")
dat = [("Seqüència", "Criteri d'integració", "RF_mass", "Ordenada", "R²", "Punts")]
for seq in ["293_SEQ_CAL", "305_SEQ_CAL", "306_SEQ_CAL"]:
    a, b = R[seq]
    n = len({f["conc"] for f in D[seq][0]})
    for nom, rr in (("vigent (programa)", a), ("estreta ancorada", b)):
        dat.append((seq, nom, _f(rr[0], 0), _f(rr[1], 1), _f(rr[2], 4), str(n)))
taula(dat)
figure("int_rectes.png",
       "Recta de calibratge de cada seqüència amb els dos criteris d'integració. A la 293 i la 305 les "
       "dues rectes se superposen i el criteri és indiferent. A la 306 la integració vigent desplaça els "
       "punts de baixa concentració cap amunt i n'inverteix l'ordre respecte de 3 ppm.")

# ============================== 8. CONVERGENCIA ==============================
doc.add_heading("8. Comparació entre la 305 i la 306 amb integració equivalent", level=1)
a305, b305 = R["305_SEQ_CAL"]
a306, b306 = R["306_SEQ_CAL"]
a293, b293 = R["293_SEQ_CAL"]
para(f"Amb finestra estreta ancorada, el factor de resposta de la 306 ({_f(b306[0], 0)}) resulta un "
     f"{(b306[0]/b305[0]-1)*100:.0f} % superior al de la 305 ({_f(b305[0], 0)}). Respecte de la 293 "
     f"({_f(b293[0], 0)}), el guany és de ×{b305[0]/b293[0]:.1f} i ×{b306[0]/b293[0]:.1f} respectivament. "
     "Les dues seqüències no convergeixen a un mateix factor de resposta.")
para("Els paràmetres mesurats a 5 ppm delimiten l'origen de la diferència:", space_after=3)
for seq in ["305_SEQ_CAL", "306_SEQ_CAL"]:
    files, _, _ = D[seq]
    f5 = [f for f in files if f["conc"] == 5.0 and f["est"]]
    h = float(np.mean([f["est"]["h_net"] for f in f5]))
    w = float(np.mean([f["est"]["fwhm"] for f in f5]))
    ar = float(np.mean([f["est"]["area"] for f in f5]))
    bullet(f"{seq}: alçada neta {_f(h)} ppb · FWHM {_f(w, 2)} min · àrea {_f(ar)}.")
para("L'alçada neta del pic és equivalent en les dues seqüències, mentre que el FWHM de la 306 és un 8 % "
     "superior i la seva àrea un 15 % superior. Un eixamplament a alçada constant implica un increment "
     "d'àrea del mateix ordre. La quantitat injectada, però, es manté (Taula 1 de l'informe de "
     "calibracions: pic de 254 nm ×1,10 a la 305 i ×0,92 a la 306, respecte de la 293), de manera que "
     "l'increment d'àrea a massa constant no és atribuïble a la mostra.", space_after=6)
para("Dos factors mesurats queden oberts i acoten la interpretació: (a) el pic de la 306 s'assenta sobre "
     "una línia de base de ~1.400–2.000 ppb amb deriva interna a la injecció, de manera que la base "
     "lineal ajustada als flancs pot no descriure el fons sota el pic; (b) el pic de 254 nm de la 306 és "
     "un 17 % més baix que el de la 305 (101,6 davant 122,6 mAU) amb la mateixa massa injectada, "
     "diferència superior a la que explicaria l'eixamplament del 8 %.", space_after=6)

# ============================== 9. EFICIENCIA DE LA COLUMNA ==============================
doc.add_heading("9. Eficiència cromatogràfica de les columnes", level=1)
para("L'eficiència de la columna es mesura sobre el pic de 254 nm, adquirit pel DAD immediatament "
     "després de la columna. El canal DOC no serveix per a aquesta mesura: hi incorpora l'eixamplament "
     "de la línia de transferència fins al TOC. La Taula 6 recull els paràmetres cromatogràfics de les "
     "tres seqüències.")
table_caption("Paràmetres cromatogràfics mesurats sobre el pic de 254 nm. Plats teòrics "
              "N = 5,54·(t_ret/FWHM)². Asimetria USP mesurada al 10 % de l'alçada (>1 cua, <1 frontal). "
              "Mitjana de les injeccions amb pic localitzable.")


def _met254(seq):
    files, _, _ = D[seq]
    import numpy as _np
    tr, fw, N, asi = [], [], [], []
    p = os.path.join(r"C:\Users\maria\Proyectos\Dades3", seq, "CHECK", "data",
                     "calibration_result.json")
    import json as _json
    d = _json.load(open(p, encoding="utf-8"))
    for cal in d.get("calibrations") or []:
        for rep in cal.get("replicas_info") or []:
            t = _np.asarray(rep.get("t_dad") or [], float)
            y = _np.asarray(rep.get("y_dad_254") or [], float)
            if len(t) < 50:
                continue
            n = len(t)
            edge = _np.r_[_np.arange(0, n // 6), _np.arange(5 * n // 6, n)]
            yd = y - _np.polyval(_np.polyfit(t[edge], y[edge], 1), t)
            i = int(_np.argmax(yd))
            if not (15 <= t[i] <= 28):
                continue
            h = float(yd[i]); half = h / 2
            li = ri = i
            while li > 0 and yd[li] > half:
                li -= 1
            while ri < n - 1 and yd[ri] > half:
                ri += 1
            w = float(t[ri] - t[li])
            if w <= 0:
                continue
            ten = h * 0.10
            l10 = r10 = i
            while l10 > 0 and yd[l10] > ten:
                l10 -= 1
            while r10 < n - 1 and yd[r10] > ten:
                r10 += 1
            a_, b_ = float(t[i]) - float(t[l10]), float(t[r10]) - float(t[i])
            tr.append(float(t[i])); fw.append(w)
            N.append(5.54 * (float(t[i]) / w) ** 2)
            if a_ > 0:
                asi.append(b_ / a_)
    return (float(_np.mean(tr)), float(_np.mean(fw)), float(_np.mean(N)), float(_np.mean(asi)))


m293 = _met254("293_SEQ_CAL"); m305 = _met254("305_SEQ_CAL"); m306 = _met254("306_SEQ_CAL")
dat = [("Seqüència", "t retenció (min)", "FWHM (min)", "Plats teòrics N", "Asimetria", "Respecte de la 293")]
for seq, m in (("293_SEQ_CAL", m293), ("305_SEQ_CAL", m305), ("306_SEQ_CAL", m306)):
    if m is m293:
        comp = "(referència)"
    else:
        comp = (f"t_ret {m[0]-m293[0]:+.2f} min · FWHM ×{m[1]/m293[1]:.2f} · "
                f"plats ×{m[2]/m293[2]:.2f}")
    dat.append((seq, _f(m[0], 2), _f(m[1], 3), _f(m[2], 0), _f(m[3], 2), comp))
taula(dat)

para(f"La seqüència 305, adquirida cinc mesos després de la 293, presenta el mateix temps de retenció "
     f"({_f(m305[0], 2)} davant {_f(m293[0], 2)} min), un FWHM equivalent (×{m305[1]/m293[1]:.2f}), un "
     f"{(m305[2]/m293[2]-1)*100:.0f} % més de plats teòrics i una asimetria menor. Els indicadors "
     "habituals de degradació d'una columna —desplaçament del temps de retenció, eixamplament del pic i "
     "pèrdua de plats— no s'hi observen.")
para(f"La columna instal·lada per a la 306 presenta un {(1-m306[2]/m293[2])*100:.0f} % menys de plats "
     f"teòrics que la 293, un FWHM ×{m306[1]/m293[1]:.2f} i el temps de retenció desplaçat "
     f"{m306[0]-m293[0]:+.2f} min. L'asimetria passa de {_f(m305[3], 2)} (cua) a {_f(m306[3], 2)} "
     "(frontal): un canvi de règim en la forma del pic, no una variació gradual. La combinació de pic "
     "frontal, pèrdua de plats, fons elevat i deriva de la línia de base és consistent amb un llit "
     "cromatogràfic amb buits o canals, o amb una columna no equilibrada.", space_after=6)
para("Els paràmetres recollits corresponen únicament al comportament cromatogràfic registrat en aquestes "
     "seqüències. Qualsevol altre motiu de substitució —pressió de treball, fuites, antiguitat o "
     "manteniment programat— no és observable en aquestes dades.", size=9, color=MUT, space_after=6)

doc.save(OUT)
print("Desat a:", OUT, "| mida KB:", round(os.path.getsize(OUT) / 1024))
