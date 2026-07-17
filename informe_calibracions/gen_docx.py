# -*- coding: utf-8 -*-
import os, sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCR=os.path.dirname(os.path.abspath(__file__))  # dades i figures a la mateixa carpeta
OUT=r"C:\Users\maria\Proyectos\Informe_calibracions_COLUMN_BP_v9.docx"

ACCENT=RGBColor(0x0E,0x62,0x6C); INK=RGBColor(0x22,0x2A,0x30); MUT=RGBColor(0x5A,0x6B,0x78)

doc=Document()
st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(10.5); st.font.color.rgb=INK
for s in ("Heading 1","Heading 2"):
    doc.styles[s].font.color.rgb=ACCENT

def para(text, size=10.5, bold=False, italic=False, color=INK, space_after=6, align=None):
    p=doc.add_paragraph(); r=p.add_run(text)
    r.font.size=Pt(size); r.bold=bold; r.italic=italic; r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(space_after)
    if align: p.alignment=align
    return p

def bullet(runs):
    p=doc.add_paragraph(style="List Bullet")
    for txt,bold in runs:
        r=p.add_run(txt); r.bold=bold; r.font.size=Pt(10.5); r.font.color.rgb=INK
    p.paragraph_format.space_after=Pt(3)
    return p

FIGN=[0]; TBLN=[0]

# Les referencies del text (Taula N / Figura N) es calculen a partir dels
# comptadors, mai a ma: inserir una taula al mig desplaca la numeracio i totes
# les referencies posteriors quedarien apuntant al lloc equivocat en silenci.
def ref_taula(n):     return f"Taula {n}"
def ref_taules(a, b): return f"Taules {a} a {b}"
def ref_figura(n):    return f"Figura {n}"
def ref_figures(a, b): return f"Figures {a} a {b}"
def seguent_taula():  return TBLN[0] + 1
def seguent_figura(): return FIGN[0] + 1

def figure(png, caption, width=6.4):
    doc.add_picture(os.path.join(SCR,png), width=Inches(width))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    FIGN[0]+=1
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f"Figura {FIGN[0]}. "); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=MUT
    r2=p.add_run(caption); r2.italic=True; r2.font.size=Pt(8.5); r2.font.color.rgb=MUT
    p.paragraph_format.space_after=Pt(12)

def table_caption(caption):
    TBLN[0]+=1
    p=doc.add_paragraph()
    r=p.add_run(f"Taula {TBLN[0]}. "); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=MUT
    r2=p.add_run(caption); r2.italic=True; r2.font.size=Pt(9); r2.font.color.rgb=MUT
    p.paragraph_format.space_after=Pt(4)
    return TBLN[0]

def shade(cell, hexcolor="0E626C"):
    tcPr=cell._tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hexcolor)
    tcPr.append(shd)

def repeat_header(table):
    trPr=table.rows[0]._tr.get_or_add_trPr()
    th=OxmlElement("w:tblHeader"); th.set(qn("w:val"),"true"); trPr.append(th)

# ---- Portada ----
t=doc.add_paragraph(); r=t.add_run("Comparació de calibracions HPSEC-DAD-DOC: modes COLUMN i BP")
r.bold=True; r.font.size=Pt(18); r.font.color.rgb=ACCENT
t.paragraph_format.space_after=Pt(4)
para("Anàlisi injecció a injecció dels patrons KHP en els canals DOC (TOC), UIB i 254 nm (DAD). "
     "Comparació entre les seqüències de calibratge anteriors i les noves, en els dos modes cromatogràfics.",
     size=11, color=MUT, space_after=10)
para("Seqüències: COLUMN 293 (20/02) → 305 (07/07) → 306 (14/07), 400 µL · BP 292 (19/02) → 304 (07/07), "
     "100 µL. Entre la 305 i la 306 s'ha substituït la columna cromatogràfica. "
     "Sensibilitat UIB 700 → 1000. Dues rèpliques per concentració.", size=9, color=MUT, space_after=14)

# ---- Resum ----
doc.add_heading("1. Resum",level=1)
para("Les seqüències de calibratge posteriors al febrer presenten un increment de la resposta del DOC en "
     "tots dos modes cromatogràfics. La incorporació de la seqüència 305 —mateix detector i mateix mode "
     "que la 306, amb una columna diferent— permet separar els dos efectes de manera directa: l'increment "
     "de resposta és present a totes tres seqüències noves (305, 306 i 304), mentre que la inestabilitat "
     "és exclusiva de la 306.")
para("Observacions principals:", bold=True, space_after=3)
bullet([("Resposta del DOC: el pic net s'incrementa un factor aproximat de 8 a totes les seqüències noves "
         "(COLUMN 305 ×7,5; COLUMN 306 ×7,6; BP 304 ×8,2, a 5 ppm). La coincidència del factor entre "
         "seqüències que comparteixen detector però no columna situa l'increment al detector i règim "
         "d'oxidació del TOC.", False)])
bullet([("L'increment de senyal NO comporta guany de sensibilitat: el soroll creix més de pressa que el "
         "senyal. La relació senyal/soroll per ppm passa de 2.329 (293) a 1.697 a la 305 (−27 %) i a 692 "
         "a la 306 (−70 %). El LOD i el LOQ empitjoren (§6).", False)])
bullet([("Línia de base (fons del DOC): ×22,7 a la 305, ×47,5 a la 306 i ×12 a la 304 (mediana del senyal "
         "cru respecte de la seqüència anterior del mateix mode).", False)])
bullet([("Estabilitat: la 305 manté una línia de base elevada però plana; la 306 presenta deriva i "
         "caigudes abruptes; la 304 (BP) manté un pedestal estable.", False)])
bullet([("254 nm: el pic es manté a la 305 (×1,10 respecte de la 293, a totes les concentracions) i a la "
         "304 (×1,09). A la 306 la línia de base del DAD deriva fins a −7 mAU i el pic queda degradat a "
         "baixa concentració (×0,58 a 0,25 ppm; ×0,92 a 5 ppm).", False)])
bullet([("UIB: a la 305 el fons és de 597 amb marge i la saturació al fons d'escala (999,63) es limita al "
         "cim del pic; a la 306 el senyal està saturat ja a la línia de base (999,6); a la 304 la "
         "saturació es restringeix a ≥3 ppm.", False)])
bullet([("Linealitat del DOC: preservada a la 305 (R² 0,998 de 0,25 a 5 ppm) i a la 304 (R² 0,999). A la "
         "306 la recta obtinguda amb la integració vigent no és lineal (R² 0,607); amb una finestra "
         "d'integració ajustada a l'amplada del pic, la linealitat es recupera (R² 0,9996). La pèrdua de "
         "linealitat de la 306 és, doncs, del criteri d'integració (§5).", False)])
bullet([("Eficiència cromatogràfica (mesurada al pic de 254 nm): la 305 manté els paràmetres de la 293 "
         "cinc mesos després (t_ret 20,89 davant 20,87 min; 25.775 davant 23.957 plats teòrics). La "
         "columna instal·lada per a la 306 en té un 19 % menys (19.319) i el pic passa d'asimetria 1,73 "
         "(cua) a 0,89 (frontal) (§4).", False)])
bullet([("Estructures posteriors al pic: els accidents a ~31 i ~36 min són presents a totes les "
         "seqüències i s'incrementen ×10,2 i ×9,1 entre la 293 i la 305, factors coincidents amb el guany "
         "del detector. L'altiplà de ~45 a ~70 min coincideix amb el tram de cabal a 1 mL/min i no depèn "
         "de la columna (§3).", False)])
bullet([("Coincidència entre injeccions: a la 305 les deu injeccions se superposen (FWHM del DOC "
         "1,407 ± 0,020 min sobre un rang de concentració de 20×) i el temps de retenció del 254 nm "
         "(20,894 ± 0,005 min) coincideix amb el de la 293 cinc mesos abans amb 0,023 min de diferència. "
         "A la 306 el 254 nm presenta una dispersió de ± 1,181 min i no és utilitzable com a referència "
         "temporal (§7).", False)])

# ---- Comparació entre seqüències (taula) ----
doc.add_heading("2. Comparació entre seqüències",level=1)
para("La Taula 1 recull la comparació de paràmetres de cada seqüència nova respecte de l'anterior del "
     "mateix mode, amb la lectura associada a cada observació. Les columnes 305 i 306 comparteixen "
     "detector, mode i volum d'injecció i difereixen únicament en la columna cromatogràfica.")
table_caption("Comparació de paràmetres de les seqüències noves respecte de l'anterior del mateix mode: "
              "COLUMN 293→305 i 293→306 (400 µL), BP 292→304 (100 µL).")
data=[
 ("Paràmetre","COLUMN (293→305)","COLUMN (293→306)","BP (292→304)","Interpretació"),
 ("Senyal 254 nm (pic)","invariant (×1,10)","×0,92 a 5 ppm; ×0,58 a 0,25 ppm","invariant (×1,09)",
  "Mateixa quantitat d'analit a 305 i 304"),
 ("Resposta DOC (pic net, 5 ppm)","×7,5","×7,6","×8,2","Comú: detector TOC"),
 ("Factor de resposta DOC (RF a 5 ppm)","324 → 3.230","324 → 3.904","56 → 435","Increment a totes tres"),
 ("Recta DOC (RF_mass, R²)","7.868 · R² 0,998","6.193 · R² 0,607","—","A la 306, efecte de la integració"),
 ("Senyal/soroll per ppm","2.329 → 1.697 (−27 %)","2.329 → 692 (−70 %)","—","El soroll creix més que el senyal"),
 ("LOD · LOQ (criteri del blanc)","×3,8 · ×3,8","×6,3 · ×6,3","—","Sensibilitat pitjor a totes dues"),
 ("LOQ empíric (rèpliques concordants)","≤ 0,25 ppm","≈ 1 ppm","—","293: ≤ 0,1 ppm"),
 ("Línia de base DOC (fons)","×22,7","×47,5","×12","Elevada a totes; màxima a la 306"),
 ("Estabilitat de la base","elevada però plana","deriva i caigudes","pedestal estable","Inestabilitat només a la 306"),
 ("Deriva de la base 254 nm","absent (~0)","severa (−7)","absent (~0)","Només a la 306"),
 ("Saturació UIB","cim del pic (fons 597)","ja a la base (999,6)","≥3 ppm","Sense marge només a la 306"),
 ("Plats teòrics (254 nm)","23.957 → 25.775 (×1,08)","23.957 → 19.319 (×0,81)","—","Columna de la 306 menys eficient"),
 ("Asimetria del pic (254 nm)","2,06 → 1,73 (cua)","2,06 → 0,89 (frontal)","—","Canvi de forma només a la 306"),
 ("Esglaó de cabal (45 min)","+34,7 % de la base","+34,2 % de la base","—","Independent de la columna; cabal +33 %"),
 ("Linealitat DOC","0,25–5 ppm (R² 0,998)","R² 0,607 (0,9996 reintegrada)","0,25–5 ppm (R² 0,999)","Preservada a 305 i 304"),
]
tbl=doc.add_table(rows=len(data),cols=5); tbl.style="Light Grid Accent 1"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,row in enumerate(data):
    for j,val in enumerate(row):
        c=tbl.cell(i,j); c.text=""
        p=c.paragraphs[0]; run=p.add_run(val); run.font.size=Pt(9)
        if i==0: run.bold=True; run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); shade(c)
doc.add_paragraph()

# ---- Efecte comú ----
doc.add_heading("3. Increment de resposta del TOC (comú a totes les seqüències noves)",level=1)
para("El factor d'increment del pic net del DOC és equivalent a les tres seqüències noves (COLUMN 305 "
     "×7,5; COLUMN 306 ×7,6; BP 304 ×8,2, a 5 ppm), malgrat diferir en columna, volum d'injecció i "
     "durada del run. A la 305 el factor es manté estable a tot el rang de concentracions (×8,1 a "
     "0,25 ppm; ×7,9 a 0,5 ppm; ×8,6 a 1 ppm; ×7,5 a 5 ppm). L'element compartit és el detector i el "
     "règim d'oxidació del TOC; per tant, l'increment és atribuïble al detector i queda confirmat de "
     "manera independent per tres sèries, dues de les quals (305 i 306) només difereixen en la columna. "
     "El canal UIB és coherent amb aquesta observació: el fons passa de 47 (293) a 597 (305) i a la "
     "saturació del fons d'escala (999,63) a la 306 (Figures 3 i 6).")

doc.add_heading("3.1. Estructures posteriors al pic i esglaó de cabal", level=2)
para("Les seqüències noves presenten, després del pic de KHP, accidents que no s'observaven anteriorment. "
     "Se'n distingeixen dos tipus, d'origen diferent.")
para("Pics discrets a ~31 i ~36 min. Són presents també a la 293, amb alçades de 8 i 42 ppb sobre la "
     "línia de base. A la 305 assoleixen 87 i 381 ppb, factors de ×10,2 i ×9,1, coincidents amb el guany "
     "del detector. No corresponen, doncs, a compostos nous: són els mateixos accidents, amplificats.",
     space_after=6)
para("Altiplà de ~45 a ~70 min. Coincideix amb el tram del mètode en què el cabal passa de 0,75 a "
     "1 mL/min. Tres observacions n'acoten l'origen: (a) apareix també a les injeccions d'aigua MQ, de "
     "manera que no prové de la mostra; (b) la seva alçada és el +34,7 % de la línia de base a la 305 i "
     "el +34,2 % a la 306 —valors coincidents malgrat tenir columnes diferents—, de manera que no depèn "
     "de la columna; (c) el +34 % coincideix amb l'increment de cabal (+33,3 %). A la 293 el mateix "
     "esglaó representa el +14,6 % de la línia de base. La resposta del detector al canvi de cabal, "
     "doncs, ha passat de la meitat de l'increment a acompanyar-lo pràcticament 1:1.", space_after=6)
table_caption("Esglaó del senyal DOC al tram de cabal de 1 mL/min (~45–70 min), mesurat sobre injeccions "
              "d'aigua MQ. Expressat com a percentatge de la línia de base de la mateixa injecció.")
d31=[("Seqüència","Línia de base (ppb)","Esglaó (% de la base)","Increment de cabal"),
     ("293_SEQ_CAL","≈ 36","+14,6 %","+33,3 %"),
     ("305_SEQ_CAL","≈ 670","+34,7 %","+33,3 %"),
     ("306_SEQ_CAL","≈ 1.500","+34,2 %","+33,3 %")]
tb31=doc.add_table(rows=len(d31),cols=4); tb31.style="Light Grid Accent 1"; tb31.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,row in enumerate(d31):
    for j,val in enumerate(row):
        c=tb31.cell(i,j); c.text=""
        run=c.paragraphs[0].add_run(val); run.font.size=Pt(9)
        if i==0: run.bold=True; run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); shade(c)
doc.add_paragraph()
para("La dispersió injecció a injecció d'aquesta mesura és ampla (del 20 % al 44 %), atès que estimar la "
     "línia de base amb el fons derivant de la 306 és sensible a la finestra escollida i les punxes de "
     "recàrrega de xeringues del TOC hi incideixen. Una injecció d'aigua MQ a cabal constant de "
     "0,75 mL/min permetria verificar aquesta observació sense aquesta incertesa.",
     size=9, color=MUT, space_after=8)

# ---- Efecte columna ----
doc.add_heading("4. Efecte associat a la columna (exclusiu de la seqüència 306)",level=1)
para("La columna cromatogràfica s'ha substituït entre la seqüència 305 i la 306. Les dues comparteixen "
     "detector, mode, volum d'injecció i mètode, de manera que la comparació aïlla l'efecte de la "
     "columna. La 305 presenta pics gaussians (R² del fit 0,994–0,999), temps de retenció constant "
     "(21,93 min a totes les injeccions), rèpliques concordants a tot el rang i línia de base del 254 nm "
     "sense deriva; la seva recta de calibratge és lineal de 0,25 a 5 ppm (R² 0,998). La 306 presenta "
     "fons ×47,5, deriva i caigudes abruptes de la línia de base, deriva del 254 nm fins a −7 mAU, UIB "
     "saturat ja a la línia de base i pics no gaussians a 0,25 i 0,5 ppm.")
para("La inestabilitat, doncs, no acompanya el canvi de columna en si sinó específicament la columna "
     "instal·lada per a la 306. El mode BP, que fa bypass de la columna, tampoc presenta aquests efectes.")

doc.add_heading("4.1. Eficiència cromatogràfica", level=2)
para("L'eficiència de la columna es mesura sobre el pic de 254 nm, adquirit pel DAD immediatament "
     "després de la columna. El canal DOC no serveix per a aquesta mesura: hi incorpora l'eixamplament "
     "de la línia de transferència fins al TOC. La Taula 3 en recull els paràmetres.")
table_caption("Paràmetres cromatogràfics mesurats sobre el pic de 254 nm. Plats teòrics "
              "N = 5,54·(t_ret/FWHM)². Asimetria USP mesurada al 10 % de l'alçada (>1 cua, <1 frontal). "
              "Mitjana de les injeccions amb pic localitzable.")
d41=[("Seqüència","t retenció (min)","FWHM (min)","Plats teòrics N","Asimetria","Respecte de la 293"),
     ("293_SEQ_CAL (20/02)","20,87","0,318","23.957","2,06","(referència)"),
     ("305_SEQ_CAL (07/07)","20,89","0,307","25.775","1,73","t_ret +0,02 · plats ×1,08"),
     ("306_SEQ_CAL (14/07)","21,93","0,372","19.319","0,89","t_ret +1,06 · plats ×0,81")]
tb41=doc.add_table(rows=len(d41),cols=6); tb41.style="Light Grid Accent 1"; tb41.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,row in enumerate(d41):
    for j,val in enumerate(row):
        c=tb41.cell(i,j); c.text=""
        run=c.paragraphs[0].add_run(val); run.font.size=Pt(9)
        if i==0: run.bold=True; run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); shade(c)
doc.add_paragraph()
para("La seqüència 305, adquirida cinc mesos després de la 293, presenta el mateix temps de retenció "
     "(20,89 davant 20,87 min), un FWHM equivalent (×0,97), un 8 % més de plats teòrics i una asimetria "
     "menor. Els indicadors habituals de degradació d'una columna —desplaçament del temps de retenció, "
     "eixamplament del pic i pèrdua de plats— no s'hi observen.")
para("La columna instal·lada per a la 306 presenta un 19 % menys de plats teòrics que la 293, un FWHM "
     "×1,17 i el temps de retenció desplaçat +1,06 min. L'asimetria passa de 1,73 (cua) a 0,89 "
     "(frontal): un canvi de règim en la forma del pic, no una variació gradual. La combinació de pic "
     "frontal, pèrdua de plats, fons elevat i deriva de la línia de base és consistent amb un llit "
     "cromatogràfic amb buits o canals, o amb una columna no equilibrada.", space_after=6)
para("Els paràmetres recollits corresponen únicament al comportament cromatogràfic registrat en aquestes "
     "seqüències. Qualsevol altre motiu de substitució —pressió de treball, fuites, antiguitat o "
     "manteniment programat— no és observable en aquestes dades.", size=9, color=MUT, space_after=8)
para("Els senyals crus dels tres canals es representen a les Figures 1 a 6: DOC continu (Figures 1 i 2), "
     "254 nm (Figures 3 i 4) i UIB (Figures 5 i 6). Les figures de COLUMN inclouen un panell per "
     "seqüència (293, 305 i 306).")

doc.add_heading("Senyal DOC continu (fulla d'adquisició 2-TOC)",level=2)
figure("raw_doc_continu.png","COLUMN, un panell per seqüència. La 293 manté la línia de base plana (~35 ppb). La 305 la té elevada (~650 ppb) però plana i sense caigudes. La 306 la té elevada (~1.400–2.000 ppb), amb deriva i caigudes abruptes. L'alçada dels pics escala amb la concentració a les tres seqüències.")
figure("bp_doc_continu.png","BP. La seqüència 304 presenta una línia de base elevada (~320 ppb) però estable; pics lineals. Sense caigudes.")

doc.add_heading("Senyal 254 nm (DAD)",level=2)
figure("raw_254_overlay.png","COLUMN, un panell per seqüència. El pic es manté a la 293 i a la 305, amb línia de base plana. A la 306 la línia de base del DAD deriva a valors negatius (fins a −7 mAU) al llarg del run.")
figure("bp_254_overlay.png","BP. Senyal 254 nm net i sense deriva en ambdues seqüències.")

doc.add_heading("Senyal UIB",level=2)
figure("raw_uib_overlay.png","COLUMN, un panell per seqüència. El fons passa de 47 (293) a 597 (305), amb marge fins al fons d'escala i saturació limitada al cim del pic. A la 306 el senyal es troba al fons d'escala (999,63) ja a la línia de base.")
figure("bp_uib_overlay.png","BP. La seqüència 304 satura únicament a 3 i 5 ppm; les concentracions baixes conserven marge.")

doc.add_heading("5. Integració del senyal DOC", level=1)
para("La finestra d'integració que obre el programa és més ampla que el pic a totes tres seqüències: "
     "mediana de 10,3 (293), 9,7 (305) i 13,0 (306) vegades el FWHM, davant les ~4 vegades que "
     "correspondrien a un pic gaussià. L'excés no té conseqüència mesurable mentre la línia de base és "
     "plana i propera a zero un cop restada: la sobre-integració resultant és de ×1,05 a ×1,44 a la 293 "
     "i de ×1,02 a ×1,20 a la 305. A la 306, amb la base elevada i derivant, la finestra incorpora "
     "senyal aliè al pic i la sobre-integració arriba a ×15,0 a 0,25 ppm.")
para("L'efecte depèn de la concentració: el senyal aliè capturat és aproximadament constant dins de cada "
     "seqüència, de manera que el seu pes relatiu creix en disminuir la concentració del patró. A la 306 "
     "això arriba a invertir l'ordre dels punts —l'àrea a 0,25 ppm (12.583) supera la de 3 ppm "
     "(11.564)— i és l'origen de la pèrdua de linealitat (R² 0,607). Reintegrant amb una finestra de "
     "±2·FWHM ancorada al pic de 254 nm, la recta de la 306 recupera la linealitat (RF_mass 9.319, "
     "ordenada −559, R² 0,9996), mentre que a la 305 el criteri d'integració és indiferent (7.868 · "
     "R² 0,998 amb el programa; 7.718 · R² 0,999 amb finestra estreta).")
para("Les àrees recollides al detall per rèplica (§9) són les que calcula el programa actualment. El "
     "detall injecció a injecció d'aquesta comparació es documenta a l'informe específic d'integració.",
     size=9, color=MUT, space_after=8)

doc.add_heading("5.1. Rectes de calibratge amb integració ajustada", level=2)
para("Quan cada injecció s'integra a la seva pròpia finestra ajustada al pic (±2·FWHM ancorada al 254 nm, "
     f"base local), les tres seqüències donen rectes de calibratge vàlides. La {ref_figura(seguent_figura())} "
     "recull els punts individuals de cada rèplica i l'ajust corresponent.")
table_caption("Recta de calibratge de cada seqüència amb totes les injeccions integrades a finestra "
              "ajustada. RF_mass = pendent (àrea per µg de DOC); n = injeccions incloses.")
d51=[("Seqüència","RF_mass","Ordenada","R²","Injeccions"),
     ("293_SEQ_CAL","757","+4,5","0,9995","12"),
     ("305_SEQ_CAL","7.718","+27,3","0,9986","10"),
     ("306_SEQ_CAL","9.270","−483,8","0,9970","9 (1 exclosa)")]
tb51=doc.add_table(rows=len(d51),cols=5); tb51.style="Light Grid Accent 1"; tb51.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,row in enumerate(d51):
    for j,val in enumerate(row):
        c=tb51.cell(i,j); c.text=""
        run=c.paragraphs[0].add_run(val); run.font.size=Pt(9)
        if i==0: run.bold=True; run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); shade(c)
doc.add_paragraph()
figure("reg_rectes.png",
       "Rectes de calibratge amb totes les injeccions integrades a finestra ajustada. Cada punt és una "
       "injecció; les dues rèpliques de cada concentració apareixen per separat. La R2 de 0,25 ppm de la "
       "306, marcada amb una creu, s'exclou perquè no conté pic de KHP.", width=6.9)
para("La 306, que amb la integració del programa dóna una recta no lineal (R² 0,607), passa a R² 0,9970 "
     "un cop integrada a finestra ajustada. La diferència, per tant, no és de les dades sinó del criteri "
     "d'integració. La R2 de 0,25 ppm queda exclosa: el pic de KHP no arriba al detector (alçada neta "
     "15 ppb davant els ~600 de la rèplica parella, SNR 7,9, i el pic de 254 nm no localitzable a la "
     "posició esperada), la qual cosa correspon a una injecció fallida i no a un problema d'integració.")
para("La R2 de 0,5 ppm de la 306 s'inclou però presenta també un dèficit d'injecció: la seva àrea és "
     "aproximadament un terç de la de la rèplica parella (alçada neta 444 davant 1.137 ppb). A la 293 i "
     "la 305 totes les rèpliques concorden. Les injeccions fallides es concentren, doncs, als punts de "
     "baixa concentració de la 306.", size=9, color=MUT, space_after=8)

doc.add_heading("6. Límits de detecció i de quantificació", level=1)
para("L'increment de la resposta del DOC no comporta un guany de sensibilitat. El soroll instrumental "
     "creix en la mateixa proporció que la línia de base —el soroll relatiu es manté entre el 0,21 % i "
     "el 0,35 % del fons a totes tres seqüències— i, en termes absoluts, més de pressa que el senyal. La "
     "relació senyal/soroll per ppm, que és qui fixa el límit de detecció, disminueix.")
table_caption("Soroll, relació senyal/soroll i límits de detecció i quantificació. LOD i LOQ en ppm de "
              "la mostra injectada (400 µL), calculats per tres criteris independents: soroll "
              "(3σ/S i 10σ/S sobre la recta d'alçades), regressió (ICH Q2: 3,3·Sy/S i 10·Sy/S) i blanc "
              "(3·SD i 10·SD de l'àrea integrada a la finestra del pic en injeccions d'aigua MQ).")
d6=[("Paràmetre","293_SEQ_CAL","305_SEQ_CAL","306_SEQ_CAL"),
    ("Línia de base (ppb)","34,9","608,0","1.485,0"),
    ("Soroll σ (ppb)","0,123","1,264","3,241"),
    ("Soroll relatiu a la base","0,35 %","0,21 %","0,22 %"),
    ("Resolució de lectura del TOC (ppb)","0,1","1,0","10,0"),
    ("Pendent (ppb/ppm)","286,5","2.144,9","2.242,2"),
    ("Senyal/soroll per ppm","2.329","1.697 (−27 %)","692 (−70 %)"),
    ("LOD · criteri del soroll (ppm)","0,0013","0,0018 (×1,4)","0,0043 (×3,4)"),
    ("LOQ · criteri del soroll (ppm)","0,0043","0,0059 (×1,4)","0,0145 (×3,4)"),
    ("LOD · criteri de la regressió (ppm)","0,147","0,262 (×1,8)","0,157 (×1,1)"),
    ("LOQ · criteri de la regressió (ppm)","0,444","0,794 (×1,8)","0,474 (×1,1)"),
    ("LOD · criteri del blanc (ppm)","0,048","0,182 (×3,8)","0,303 (×6,3)"),
    ("LOQ · criteri del blanc (ppm)","0,161","0,606 (×3,8)","1,009 (×6,3)"),
    ("LOQ empíric (rèpliques concordants)","≤ 0,1 ppm","≤ 0,25 ppm","≈ 1 ppm")]
tb6=doc.add_table(rows=len(d6),cols=4); tb6.style="Light Grid Accent 1"; tb6.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,row in enumerate(d6):
    for j,val in enumerate(row):
        c=tb6.cell(i,j); c.text=""
        run=c.paragraphs[0].add_run(val); run.font.size=Pt(9)
        if i==0: run.bold=True; run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); shade(c)
doc.add_paragraph()
para("Tres observacions acoten la lectura de la taula:", bold=True, space_after=3)
bullet([("El criteri de la regressió (ICH Q2) opera sobre les mitjanes de rèpliques i no detecta la "
         "discordança entre elles. A la 306 assigna un LOD de 0,157 ppm, inferior al de la 305, tot i "
         "que a 0,25 ppm les dues rèpliques difereixen un 1.034 % (alçades netes de 617 i 15 ppb) i a "
         "0,5 ppm un 62 %. El criteri del blanc i el LOQ empíric sí que ho recullen.", False)])
bullet([("La resolució de lectura del TOC (tres xifres significatives) creix amb el nivell de senyal: "
         "0,1 ppb a la 293, 1 ppb a la 305 i 10 ppb a la 306. A la 306 aquests 10 ppb equivalen a "
         "0,0045 ppm, valor coincident amb el LOD per soroll (0,0043 ppm): el límit hi queda determinat "
         "per la quantització de l'instrument i no pel soroll.", False)])
bullet([("El LOQ empíric —concentració més baixa on les rèpliques concorden amb un RSD ≤ 10 %— és de "
         "≤ 0,1 ppm a la 293 (RSD 0,1 %), ≤ 0,25 ppm a la 305 (RSD 1,7 %; concentració més baixa "
         "assajada) i ≈ 1 ppm a la 306 (RSD 1,6 %, davant 1.034 % a 0,25 ppm i 62 % a 0,5 ppm).", False)])

doc.add_heading("7. Solapament dels pics", level=1)
para(f"Les {ref_figures(seguent_figura(), seguent_figura()+1)} superposen els pics de totes les "
     "injeccions de cada seqüència, en cru i normalitzats a l'alçada. És la prova directa de "
     "coincidència: si la columna i l'adquisició són estables, totes les injeccions han d'eluir al "
     "mateix temps i, un cop normalitzades, han de tenir la mateixa forma amb independència de la "
     "concentració.")
figure("sol_doc_solapat.png",
       "Canal DOC. Fila superior: pics crus nets de base local. Fila inferior: els mateixos pics "
       "normalitzats a l'alçada i referits al temps de retenció mitjà de la seqüència. A la 305 les deu "
       "injeccions se superposen en una sola corba al llarg d'un rang de concentració de 20×. A la 306, "
       "les injeccions de 0,25 ppm mostren l'escalonament de la quantització del TOC (10 ppb sobre un pic "
       "de 15 ppb) i una de 0,5 ppm apareix desplaçada.", width=6.9)
figure("sol_254_solapat.png",
       "Canal 254 nm, mateixa disposició. A la 293 i la 305 tots els pics coincideixen. El màxim "
       "secundari a +0,57 min no és una espatlla del pic de KHP sinó un pic resolt de matèria diferent "
       "(§8): la seva alçada absoluta és constant i per això només és apreciable a les concentracions "
       "baixes. A la 306 les concentracions de 0,25 i 0,5 ppm no permeten localitzar el pic amb "
       "fiabilitat a causa de la deriva de la línia de base del DAD.", width=6.9)

table_caption("Coincidència entre injeccions dins de cada seqüència: temps de retenció i FWHM del pic, "
              "mitjana i dispersió (desviació estàndard) sobre totes les injeccions de la seqüència. "
              "Canals DOC i 254 nm.")
d71=[("Seqüència","Canal","t retenció (min)","Dispersió","Rang","FWHM (min)","Dispersió"),
     ("293_SEQ_CAL","DOC","22,722","± 0,046","0,133","1,119","± 0,094"),
     ("305_SEQ_CAL","DOC","21,907","± 0,037","0,100","1,407","± 0,020"),
     ("306_SEQ_CAL","DOC","23,843","± 0,116","0,400","1,520","± 0,221"),
     ("293_SEQ_CAL","254 nm","20,871","± 0,007","0,015","0,318","± 0,009"),
     ("305_SEQ_CAL","254 nm","20,894","± 0,005","0,014","0,308","± 0,009"),
     ("306_SEQ_CAL","254 nm","21,540","± 1,181","4,213","0,435","± 0,184")]
tb71=doc.add_table(rows=len(d71),cols=7); tb71.style="Light Grid Accent 1"; tb71.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,row in enumerate(d71):
    for j,val in enumerate(row):
        c=tb71.cell(i,j); c.text=""
        run=c.paragraphs[0].add_run(val); run.font.size=Pt(9)
        if i==0: run.bold=True; run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); shade(c)
doc.add_paragraph()
para("El temps de retenció del 254 nm de la 305 (20,894 ± 0,005 min) coincideix amb el de la 293 "
     "(20,871 ± 0,007 min) amb una diferència de 0,023 min —un segon i mig— cinc mesos després. El FWHM "
     "del DOC de la 305 és de 1,407 ± 0,020 min sobre les deu injeccions, amb un rang de 0,067 min. A la "
     "306 el temps de retenció del 254 nm presenta una dispersió de ± 1,181 min i un rang de 4,213 min, "
     "de manera que el canal no és utilitzable com a referència temporal en aquesta seqüència.")
para("El FWHM es manté constant amb la concentració a la 293 i la 305 en tots dos canals (quocient entre "
     "el punt de 5 ppm i el més baix: 0,79 i 1,02 al DOC; 0,92 i 0,96 al 254 nm), de manera que no "
     "s'observa sobrecàrrega de la columna en cap de les dues. A la 306 el FWHM del 254 nm a 0,25 ppm "
     "(0,687 min) és el doble del de 5 ppm (0,367 min).", space_after=8)

import json
def _f(v,dec):
    if not isinstance(v,(int,float)): return "—"
    return f"{v:,.{dec}f}".replace(",","§").replace(".",",").replace("§",".")

def _uib_cell(r, force_sat):
    if force_sat: return "sat"
    ym=r.get("uib_ymax")
    if isinstance(ym,(int,float)) and ym>=999: return "sat"
    return _f(r.get("uib_area"),1)

def inj_table(rows, caption, force_sat=False):
    table_caption(caption)
    hdr=("KHP","ppm","Rep","DOC àrea","DOC R²","UIB àrea","254 àrea")
    rows=sorted(rows, key=lambda x:(x["conc"], str(x["rep"])))
    tb=doc.add_table(rows=len(rows)+1, cols=len(hdr)); tb.style="Light Grid Accent 1"
    for j,h in enumerate(hdr):
        run=tb.cell(0,j).paragraphs[0].add_run(h); run.bold=True; run.font.size=Pt(8.5); run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        shade(tb.cell(0,j))
    repeat_header(tb)
    for i,r in enumerate(rows,1):
        anom = r.get("doc_anoms") or []
        bad = "KHP_PEAK_NON_GAUSSIAN" in anom
        uib = _uib_cell(r, force_sat)
        vals=(r["khp"], str(r["conc"]).replace(".",","), "R"+str(r["rep"]),
              _f(r.get("doc_area"),1), _f(r.get("doc_r2"),3), uib, _f(r.get("a254_area"),2))
        for j,v in enumerate(vals):
            run=tb.cell(i,j).paragraphs[0].add_run(v); run.font.size=Pt(8.5)
            if bad and j==3: run.font.color.rgb=RGBColor(0xBE,0x3A,0x38); run.bold=True
            if v=="sat": run.font.color.rgb=RGBColor(0xBE,0x3A,0x38)
    doc.add_paragraph()

doc.add_heading("8. Espectres UV dels pics d'una injecció de KHP", level=1)
para("El detector DAD adquireix l'espectre complet (200–400 nm, 101 longituds d'ona) a cada punt del "
     "cromatograma, de manera que cada pic es pot caracteritzar espectralment sense injectar patrons "
     f"addicionals. La {ref_figura(seguent_figura())} recull el cromatograma d'una injecció de KHP de "
     f"5 ppm de la 293 a diverses longituds d'ona i la {ref_figura(seguent_figura()+1)} els espectres "
     "dels pics que s'hi localitzen.")
para("Criteris de lectura: el KHP és hidrogenftalat de potassi, un compost aromàtic, i presenta per tant "
     "la banda del benzè amb cua fins a ~280 nm a més de l'absorció intensa per sota de 220 nm. Les "
     "espècies inorgàniques que absorbeixen a l'UV (nitrat, nitrit) tenen una banda intensa per sota de "
     "230 nm i són transparents per sobre de 240 nm: no presenten banda aromàtica.", space_after=8)

figure("esp_cromatograma.png",
       "293_SEQ_CAL, KHP 5 ppm rèplica 1. Cromatograma del DAD a 210, 230, 254 i 280 nm, amb la línia de "
       "base mòbil llevada. Panell inferior: mateix cromatograma amb l'escala ampliada. Línies de punts: "
       "posició dels pics localitzats. A 210 nm es fan visibles estructures que a 254 nm són "
       "indistingibles del soroll.", width=6.9)

N_ESP = table_caption(
    "Paràmetres espectrals dels pics localitzats en una injecció de KHP de 5 ppm "
    "(293_SEQ_CAL). A254/Amàx i A280/Amàx són l'absorbància relativa al màxim de l'espectre "
    "de cada pic; A280/Amàx quantifica el caràcter aromàtic.")
d8=[("t (min)","A a 210 nm","A a 254 nm","λmàx (nm)","A254/Amàx","A280/Amàx","Lectura"),
    ("20,86 — KHP","1.182,8","109,8","200","0,048","0,022","Banda aromàtica present"),
    ("21,50","131,0","1,96","206","0,014","0,008","Sense banda aromàtica"),
    ("25,40","4,8","0,03","200","−0,001","0,003","Sense banda aromàtica"),
    ("36,89","3,0","0,09","200","0,008","−0,001","Sense banda aromàtica")]
tb8=doc.add_table(rows=len(d8),cols=7); tb8.style="Light Grid Accent 1"; tb8.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,row in enumerate(d8):
    for j,val in enumerate(row):
        c=tb8.cell(i,j); c.text=""
        run=c.paragraphs[0].add_run(val); run.font.size=Pt(9)
        if i==0: run.bold=True; run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); shade(c)
doc.add_paragraph()

figure("esp_espectres.png",
       f"Espectres UV dels pics de la {ref_taula(N_ESP)}, crus i normalitzats al seu màxim. El pic de "
       "KHP (20,86 min) és l'únic que presenta la cua aromàtica entre 260 i 285 nm. Els altres tres "
       "tenen el màxim per sota de 210 nm i són transparents per sobre de 240 nm.")

para("El pic de 21,50 min (a 21,45 min a les injeccions de concentració baixa, on queda resolt) mereix "
     "atenció particular per tres motius:", bold=True, space_after=3)
bullet([("És present a totes les injeccions, incloses les d'aigua MQ: a la 293 s'hi mesura a 21,45 min "
         "amb alçades de 0,72 a 1,04 mAU i a la 305 amb alçades de 0,58 a 1,08 mAU. No prové, doncs, de "
         "la mostra.", False)])
bullet([("La seva alçada és independent de la concentració de KHP: 0,94 · 0,93 · 0,95 · 0,95 mAU a 0,1 i "
         "0,25 ppm de la 293 (CV 1 % amb 2,5 vegades més patró). No forma part del KHP.", False)])
bullet([("El seu espectre té el màxim a 206 nm i és pla a partir de 240 nm. El senyal que se n'observa a "
         "254 nm (0,9 mAU) és la cua d'aquesta banda, no un cromòfor propi: a 206 nm el pic assoleix "
         "101 mAU, valor comparable al del propi KHP a 0,1 ppm (89 mAU a 200 nm).", False)])
para("Al canal DOC, al temps que li correspon un cop aplicat el desfasament DOC−254 de la seqüència "
     "(23,30 min a la 293), les injeccions d'aigua MQ presenten entre +2 i +3 ppb sobre la línia de base, "
     "davant els ~11 ppb que correspondrien si tingués la mateixa relació carboni/absorbància que el KHP. "
     "Aquesta mesura, però, queda dins la cua d'una estructura pròpia del blanc centrada a 21,6 min, de "
     "manera que la xifra és un límit superior.")
para("Absorció intensa a l'UV llunyà, absència de banda aromàtica i contingut de carboni molt inferior al "
     "que correspondria a la seva absorbància són compatibles amb un anió inorgànic; el nitrat presenta "
     "un màxim a 201–210 nm i transparència per sobre de 240 nm. La identificació requeriria injectar el "
     "patró corresponent.", space_after=6)
para("Conseqüència sobre les mesures: a 5 ppm el pic representa el 1,8 % de l'alçada del KHP a 254 nm i "
     "és negligible; a 0,1 ppm n'és el 38 % i a 0,25 ppm el 29 %, de manera que l'àrea de 254 nm dels "
     "punts baixos de la recta queda sobrevalorada. A 210 o 220 nm el pic és del mateix ordre que el "
     "KHP.", size=9, color=MUT, space_after=8)

doc.add_heading("9. Detall injecció a injecció",level=1)
para(f"El detall per rèplica dels tres canals es recull a les "
     f"{ref_taules(seguent_taula(), seguent_taula()+4)} (una per seqüència). Permet localitzar rèpliques "
     "anòmales i verificar la linealitat punt a punt. Àrees en unitats d'integració; 'sat' indica "
     "saturació de l'UIB al fons d'escala (999,63).", color=MUT, space_after=8)
col=json.load(open(SCR+r"\cal_compare_data.json",encoding="utf-8"))
bp=json.load(open(SCR+r"\bp_data.json",encoding="utf-8"))
doc.add_heading("Mode COLUMN",level=2)
inj_table(col["293_SEQ_CAL"]["rows"], "COLUMN, seqüència 293_SEQ_CAL (anterior; UIB sens. 700, 400 µL). Àrees dels canals DOC, UIB i 254 nm per rèplica de KHP.")
inj_table(col["305_SEQ_CAL"]["rows"], "COLUMN, seqüència 305_SEQ_CAL (nova, columna anterior a la substitució; UIB sens. 1000, 400 µL). Àrees dels canals DOC, UIB i 254 nm per rèplica de KHP. L'UIB assoleix el fons d'escala al cim del pic a totes les injeccions.", force_sat=True)
inj_table(col["306_SEQ_CAL"]["rows"], "COLUMN, seqüència 306_SEQ_CAL (nova, columna substituïda; UIB sens. 1000, 400 µL). Àrees dels canals DOC, UIB i 254 nm per rèplica de KHP. L'UIB es troba al fons d'escala ja a la línia de base.", force_sat=True)
doc.add_heading("Mode BP",level=2)
inj_table(bp["292_SEQ_CAL_BP"]["rows"], "BP, seqüència 292_SEQ_CAL_BP (anterior; UIB sens. 700, 100 µL). Àrees dels canals DOC, UIB i 254 nm per rèplica de KHP.")
inj_table(bp["304_SEQ_CAL_BP"]["rows"], "BP, seqüència 304_SEQ_CAL_BP (nova; UIB sens. 1000, 100 µL). Àrees dels canals DOC, UIB i 254 nm per rèplica de KHP.")

doc.save(OUT)
print("Desat a:", OUT, "| mida KB:", round(os.path.getsize(OUT)/1024))
